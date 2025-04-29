import PyPDF2
import re, json
import os
import docx
import pandas as pd
import mimetypes
from random import randint
from pydantic import BaseModel,Field
from crewai.flow import Flow, listen, start
from typing import Dict, Any, Optional, List
from regulatory_compliance.crews.document_analyzer.document_analyze_crew import DocumentAnalyzer
from regulatory_compliance.crews.logistic_compliant.logistic_compliant_crew import LogisticCompliant
from regulatory_compliance.crews.technical_validator.technical_validator_crew import TechnicalValidator
from regulatory_compliance.crews.report_generator.report_generator_crew import ReportGenerator
import html
from docx.shared import RGBColor
from io import BytesIO
import shutil
import fitz  

from dotenv import load_dotenv

load_dotenv()

def json_to_markdown(data, level=1):
    markdown = ""
    
    if isinstance(data, dict):
        for key, value in data.items():
            markdown += f"**{key.replace('_', ' ').title()}**\n\n"
            markdown += json_to_markdown(value, level + 1) + "\n"
    elif isinstance(data, list):
        for item in data:
            markdown += f"- {json_to_markdown(item, level + 1).strip()}\n"
    else:
        markdown += f"{data}\n"

    return markdown

class ComplianceState(BaseModel):
    document: str = Field(default="", description="User uploaded document")
    document_analysis: Dict[str, Any] = Field(default_factory=dict, description="analysis of the document")
    output_path: str = "compliance_report.md"
    original_file_path: str = Field(default="", description="Path to the original document")
    document_format: str = Field(default="", description="Format of the uploaded document")
    country: str = Field(default="", description="Country for regulatory compliance checks")


class ComplianceFlow(Flow[ComplianceState]):
  
    @start()
    def select_country(self):
        print("\n Logistics Regulatory Compliance Checker ")
        print("\nFirst, please specify which country's regulations to check against.")
        
        # List of supported countries
        supported_countries = [
            "United States", "European Union", "United Kingdom", "Canada", 
            "Australia", "China", "Japan", "India", "Brazil", "Germany",
            "France", "Mexico", "Russia", "South Africa", "Singapore"
        ]
        
        # Display options
        print("\nSupported countries:")
        for i, country in enumerate(supported_countries, 1):
            print(f"{i}. {country}")
        
        # Custom option
        print(f"{len(supported_countries) + 1}. Other (specify)")
        
        while True:
            try:
                choice = input("\nEnter the number of your selection: ")
                if not choice.strip():
                    raise ValueError("Please enter a selection")
                    
                choice = int(choice)
                if 1 <= choice <= len(supported_countries):
                    self.state.country = supported_countries[choice - 1]
                    break
                elif choice == len(supported_countries) + 1:
                    custom_country = input("Enter the country name: ")
                    if custom_country.strip():
                        self.state.country = custom_country
                        break
                    else:
                        print("Country name cannot be empty.")
                else:
                    print(f"Please enter a number between 1 and {len(supported_countries) + 1}")
            except ValueError:
                print("Please enter a valid number")
        
        print(f"\nSelected country for compliance checks: {self.state.country}")
        return self.state
    
    @listen(select_country)
    def get_user_input_and_extract_text(self, state):
        print("\nNow, please provide the document to analyze.")

        while True:
            file_path = input("Enter the path to the document file to analyze: ")
            try:
                if not os.path.exists(file_path):
                    raise FileNotFoundError

                mime_type, _ = mimetypes.guess_type(file_path)
                text = ""
                document_format = ""

                if file_path.lower().endswith(".pdf"):
                    with open(file_path, "rb") as file:
                        reader = PyPDF2.PdfReader(file)
                        for page in reader.pages:
                            text += page.extract_text() or ""
                    document_format = "pdf"

                elif file_path.lower().endswith(".docx"):
                    doc = docx.Document(file_path)
                    for para in doc.paragraphs:
                        text += para.text + "\n"
                    document_format = "docx"

                elif file_path.lower().endswith(".txt"):
                    with open(file_path, "r", encoding="utf-8") as file:
                        text = file.read()
                    document_format = "txt"

                # elif file_path.lower().endswith(".xlsx"):
                #     df = pd.read_excel(file_path)
                #     text = df.to_string(index=False)
                #     document_format = "xlsx"

                else:
                    print("Unsupported file format. Please try PDF, DOCX or TXT.")
                    continue

                self.state.document += text
                self.state.original_file_path = file_path
                self.state.document_format = document_format
                break

            except FileNotFoundError:
                print(f"Error: File '{file_path}' not found. Please try again.")
            except Exception as e:
                print(f"Error reading file: {str(e)}. Please try again.")

        return self.state
    
    @listen(get_user_input_and_extract_text)
    def analyze_document(self, state):
        print("Step 1: Analyzing the document ...")
        
        document_analyzer = DocumentAnalyzer()
        result = document_analyzer.crew().kickoff(
            inputs={
                "document": state.document,
                "country": state.country
            }
        )
        
        json_result = re.search(r'```json(.*)```', result.raw, re.DOTALL)
        if json_result:
            document_analysis = json.loads(json_result.group(1))
        else:
            document_analysis = json.loads(result.raw)
            markdown_output = json_to_markdown(document_analysis)
        
        self.state.document_analysis = document_analysis
        
        print(f"Document analysis complete. Identified {len(document_analysis.get('applicable_regulations', []))} applicable regulations for {state.country}.")
        return self.state.document_analysis

   
    @listen(analyze_document)
    def evaluate_regulatory_compliance(self, document_analysis):
        print("Step 2: Evaluating regulatory compliance ...")
        
        logistic_compliance_crew = LogisticCompliant()
        result = logistic_compliance_crew.crew().kickoff(
            inputs={
                "document_analysis": json.dumps(document_analysis),
                "country": self.state.country
            }
        )
        
        json_result = re.search(r'```json(.*)```', result.raw, re.DOTALL)
        if json_result:
            regulatory_compliance_evaluation = json.loads(json_result.group(1))
        else:
            regulatory_compliance_evaluation = json.loads(result.raw)
            markdown_output = json_to_markdown(regulatory_compliance_evaluation)
        
        # Check if non-compliant sections are found and log them
        if "non_compliant_sections" in regulatory_compliance_evaluation:
            non_compliant_count = len(regulatory_compliance_evaluation["non_compliant_sections"])
            print(f"Found {non_compliant_count} non-compliant sections in the document based on {self.state.country} regulations.")
        else:
            print(f"No specific non-compliant sections identified in the document based on {self.state.country} regulations.")
        
        evaluation_data={  
            "regulatory_compliance_evaluation": regulatory_compliance_evaluation,
            "document_analysis": document_analysis,
            "country": self.state.country
        }
        print(f"Regulatory compliance evaluation for {self.state.country} complete.")
        return evaluation_data
    
    @listen(evaluate_regulatory_compliance)
    def assess_technical_controls(self, data):
        print("Step 3: Assessing technical controls...")
        technical_validation_crew = TechnicalValidator()
        result = technical_validation_crew.crew().kickoff(
            inputs={
                "document_analysis": json.dumps(data['document_analysis']),
                "regulatory_compliance_evaluation": json.dumps(data['regulatory_compliance_evaluation']),
                "country": data['country']
            }
        )
        
        # Parse the JSON result
        try:
            json_result = re.search(r'```json(.*)```', result.raw, re.DOTALL)
            if json_result:
                technical_assessment = json.loads(json_result.group(1))
            else:
                # Check if result.raw is valid JSON
                if result.raw and result.raw.strip():
                    technical_assessment = json.loads(result.raw)
                else:
                    # Handle empty or invalid JSON result
                    print("Warning: Received empty or invalid response from technical validator")
                    technical_assessment = {"error": "Technical validation failed or returned empty result"}
                    
            markdown_output = json_to_markdown(technical_assessment)
        except json.JSONDecodeError as e:
            print(f"Error parsing technical assessment: {str(e)}")
            # Create fallback assessment with error info
            technical_assessment = {
                "error": f"Failed to parse technical validation results: {str(e)}",
                "raw_data": result.raw[:100] + "..." if len(result.raw) > 100 else result.raw
            }
            
        technical_assessment_data={
            "technical_assessment": technical_assessment,
            "regulatory_compliance_evaluation": data['regulatory_compliance_evaluation'],
            "document_analysis": data['document_analysis'],
            "country": data['country']
        }
        print(f"Technical assessment for {data['country']} complete.")
        return technical_assessment_data
    
    @listen(assess_technical_controls)
    def generate_compliance_report(self, data):
        """Generate a comprehensive compliance report using the Report Generator Crew"""
        print("Step 4: Generating compliance report...")
        
        # Run the report generator crew
        report_generator = ReportGenerator()
        result = report_generator.crew().kickoff(
            inputs={
                "document_analysis": json.dumps(data['document_analysis']),
                "regulatory_compliance_evaluation": json.dumps(data['regulatory_compliance_evaluation']),
                "technical_assessment": json.dumps(data['technical_assessment']),
                "country": data['country']
            }
        )
        
        try:
            report = re.search(r'```markdown(.*)```', result.raw, re.DOTALL)
            if report:
                compliance_report = report.group(1)
            else:
                compliance_report = result.raw  
        except Exception as e:
            print(f"Error parsing JSON result: {str(e)}")
            compliance_report = result.raw
        
        # Add highlighted non-compliant sections if available
        non_compliant_sections = []
        try:
            if "non_compliant_sections" in data["regulatory_compliance_evaluation"]:
                non_compliant_sections = data["regulatory_compliance_evaluation"]["non_compliant_sections"]
                highlighted_sections_markdown = f"\n\n## Highlighted Non-Compliant Sections for {data['country']}\n\n"
                
                for section in non_compliant_sections:
                    print(f"- Header: '{section.get('section_header', 'NONE')}'")
                    print(f"  Text: '{section.get('text', 'NONE')}'")   
                    highlighted_sections_markdown += f"### Non-Compliant Text\n\n"
                    highlighted_sections_markdown += f"```diff\n- {section['text']}\n```\n\n"
                    highlighted_sections_markdown += f"**Violates Regulation:** {section['regulation']}\n\n"
                    highlighted_sections_markdown += f"**Issue:** {section['issue']}\n\n"
                    highlighted_sections_markdown += f"**Recommendation:** {section['recommendation']}\n\n"
                    highlighted_sections_markdown += "---\n\n"
                
                compliance_report += highlighted_sections_markdown
        except Exception as e:
            print(f"Error adding highlighted non-compliant sections: {str(e)}")
        
        # Add country information to the report title
        if isinstance(compliance_report, str):
            country_title = f"# Regulatory Compliance Report for {data['country']}\n\n"
            if compliance_report.startswith("# "):
                # Replace the existing title
                first_newline = compliance_report.find("\n")
                compliance_report = country_title + compliance_report[first_newline+1:]
            else:
                # Add title at the beginning
                compliance_report = country_title + compliance_report
        
        # Save the markdown report to a file
        self._save_markdown_report(compliance_report, self.state.output_path)
        
        # Create highlighted version of the original document
        if non_compliant_sections and self.state.original_file_path:
            try:
                self._create_highlighted_document(non_compliant_sections)
            except Exception as e:
                print(f"Error creating highlighted document: {str(e)}")
        
        # Extract compliance status if available
        try:
            # Get compliance status from regulatory compliance evaluation
            if "regulatory_compliance_evaluation" in data and "compliance_status" in data["regulatory_compliance_evaluation"]:
                compliance_status = data["regulatory_compliance_evaluation"]["compliance_status"]
                # Save compliance status to a separate JSON file
                compliance_status_path = os.path.splitext(self.state.output_path)[0] + "_status.json"
                with open(compliance_status_path, "w") as f:
                    json.dump(compliance_status, f, indent=2)
                print(f"Compliance status saved to {compliance_status_path}")
                
                # Print compliance status summary
                print(f"\nCOMPLIANCE STATUS SUMMARY FOR {data['country']}:")
                for regulation, status in compliance_status.items():
                    print(f"\n{regulation}:")
                    print(f"Status: {status['status']}")
                    if status['issues']:
                        print("Issues:")
                        for issue in status['issues']:
                            print(f"- {issue}")
                    if status['recommendations']:
                        print("Recommendations:")
                        for rec in status['recommendations']:
                            print(f"- {rec}")
        except Exception as e:
            print(f"Error extracting compliance status: {str(e)}")
        
        # Print summary of non-compliant sections
        try:
            if non_compliant_sections:
                print(f"\nNON-COMPLIANT SECTIONS FOUND (based on {data['country']} regulations):")
                
                for i, section in enumerate(non_compliant_sections, 1):
                    print(f"\n{i}. Non-Compliant Text: {section['text'][:50]}...")
                    print(f"   Violates: {section['regulation']}")
                    print(f"   Issue: {section['issue']}")
        except Exception as e:
            print(f"Error printing non-compliant sections summary: {str(e)}")
        
        print(f"\nCompliance report for {data['country']} generated and saved to {self.state.output_path}")
        
        # Print summary information
        if isinstance(compliance_report, dict) and "executive_summary" in compliance_report:
            print("\nEXECUTIVE SUMMARY:")
            print(compliance_report["executive_summary"])
        
        if isinstance(compliance_report, dict) and "prioritized_actions" in compliance_report:
            print("\nPRIORITIZED ACTIONS:")
            for i, action in enumerate(compliance_report["prioritized_actions"], 1):
                print(f"{i}. {action['description']} (Priority: {action['priority']})")
        
        return compliance_report

    def _save_markdown_report(self, markdown_content: str, output_path: str) -> str:
        """Save the markdown report to a file with highlighted empty fields"""
        # Create the directory if it doesn't exist
        os.makedirs(os.path.dirname(output_path) or '.', exist_ok=True)
        
        # Get list of missing fields from regulatory compliance evaluation
        empty_fields = []
        non_compliant_text = []
        
        try:
            if hasattr(self.state, 'document_analysis') and isinstance(self.state.document_analysis, dict):
                non_compliant_sections = self.state.document_analysis.get("regulatory_compliance_evaluation", {}).get("non_compliant_sections", [])
                
                # Dynamically extract fields based on the compliance issues
                for section in non_compliant_sections:
                    section_header = section.get('section_header', '')
                    issue = section.get('issue', '').lower()
                    text = section.get('text', '').lower()
                    
                    # Determine if this is a missing field issue
                    if ('missing' in issue or 'empty' in issue or 'required' in issue or 
                        'missing' in text or 'empty' in text or 'not provided' in text):
                        if section_header:
                            empty_fields.append(section_header)
                    elif text and section.get('text', ''):  # Non-compliant text that's not missing
                        non_compliant_text.append(section.get('text', ''))
        except Exception as e:
            print(f"Warning: Error extracting data for report highlighting: {str(e)}")
        
        # Highlight the empty fields in the markdown content
        updated_markdown = markdown_content
        
        # First highlight empty fields
        for field in empty_fields:
            # Skip empty field names
            if not field:
                continue
                
            # Create HTML-style highlighted version for the markdown
            highlighted_field = f"**<span style='color:red'>{field}</span>**"
            
            # Replace the field with highlighted version in markdown content
            # Be careful to only replace exact matches to avoid partial replacements
            pattern = re.compile(r'\b' + re.escape(field) + r'\b')
            updated_markdown = pattern.sub(highlighted_field, updated_markdown)
        
        # Then highlight non-compliant text
        for text in non_compliant_text:
            if not text or len(text) < 5:  # Skip very short text
                continue
                
            # Create HTML-style highlighted version for the markdown
            highlighted_text = f"<span style='color:red'>{text}</span>"
            
            # Replace the text with highlighted version in markdown
            if text in updated_markdown:
                updated_markdown = updated_markdown.replace(text, highlighted_text)
        
        # Add a section about missing fields if found
        if empty_fields:
            empty_fields_section = "\n\n## Missing Required Fields\n\n"
            empty_fields_section += "The following fields are required but missing from the document:\n\n"
            
            for field in empty_fields:
                if field:
                    empty_fields_section += f"- **<span style='color:red'>{field}</span>**: Required field is empty\n"
            
            # Add the section before any closing tags if present, otherwise at the end
            if "</body>" in updated_markdown:
                updated_markdown = updated_markdown.replace("</body>", empty_fields_section + "\n</body>")
            else:
                updated_markdown += empty_fields_section
        
        # Add CSS styling if not already present
        if "<style>" not in updated_markdown:
            css_style = "\n\n<style>\n"
            css_style += "span.missing { color: red; font-weight: bold; }\n"
            css_style += "span.non-compliant { color: red; }\n"
            css_style += "</style>\n\n"
            
            # Insert after first heading or at the beginning
            if "# " in updated_markdown:
                first_heading_end = updated_markdown.find("# ") + updated_markdown[updated_markdown.find("# "):].find("\n")
                updated_markdown = updated_markdown[:first_heading_end+1] + css_style + updated_markdown[first_heading_end+1:]
            else:
                updated_markdown = css_style + updated_markdown
        
        with open(output_path, "w") as f:
            f.write(updated_markdown)
            
        return output_path
    
    def _create_highlighted_document(self, non_compliant_sections):
        """Create a highlighted version of the original document with non-compliant sections in red"""
        original_path = self.state.original_file_path
        doc_format = self.state.document_format
        
        # Create output path for the highlighted document
        filename, ext = os.path.splitext(original_path)
        highlighted_path = f"{filename}_highlighted{ext}"
        
        print(f"Creating highlighted document at: {highlighted_path}")
        
        # Handle different document formats
        if doc_format == "pdf":
            self._highlight_pdf(original_path, highlighted_path, non_compliant_sections)
        elif doc_format == "docx":
            self._highlight_docx(original_path, highlighted_path, non_compliant_sections)
        elif doc_format == "txt":
            self._highlight_txt(original_path, highlighted_path, non_compliant_sections)
        elif doc_format == "xlsx":
            print("Warning: Highlighting Excel files is not supported. Skipping highlight generation.")
        else:
            print(f"Warning: Highlighting for format {doc_format} is not supported.")
        
        return highlighted_path

    def _highlight_pdf(self, original_path, highlighted_path, non_compliant_sections):
        """Highlight non-compliant sections in a PDF document dynamically without hardcoded fields"""
        try:
            # Copy the original PDF
            shutil.copy(original_path, highlighted_path)
            
            # Open the PDF
            doc = fitz.open(highlighted_path)
            highlight_count = 0
            
            print(f"\nHighlighting non-compliant sections in PDF...")
            
            # First pass: Extract all fields that need highlighting from the non_compliant_sections data
            fields_to_highlight = {}
            
            for section in non_compliant_sections:
                section_header = section.get('section_header', '')
                text = section.get('text', '').strip()
                issue = section.get('issue', '')
                regulation = section.get('regulation', '')
                recommendation = section.get('recommendation', '')
                
                # Determine if this is a missing field issue
                is_missing_field = False
                if ('missing' in issue.lower() or 'empty' in issue.lower() or 'required' in issue.lower() or 
                    'missing' in text.lower() or 'empty' in text.lower() or 'not provided' in text.lower()):
                    is_missing_field = True
                
                # Add to the appropriate category
                if is_missing_field and section_header:
                    fields_to_highlight[section_header] = {
                        'type': 'missing_field',
                        'data': section
                    }
                elif text and not is_missing_field:
                    fields_to_highlight[text] = {
                        'type': 'non_compliant_text',
                        'data': section
                    }
            
            # Second pass: Find positions of all fields and text to highlight
            field_positions = {}
            for page_num in range(len(doc)):
                page = doc[page_num]
                page_text = page.get_text()
                
                # Check for each field or text to highlight
                for field, info in fields_to_highlight.items():
                    if field in page_text:
                        pos = page_text.find(field)
                        field_positions[field] = {
                            'page': page_num,
                            'pos': pos,
                            'type': info['type'],
                            'data': info['data']
                        }
                        print(f"Found {info['type']} '{field}' on page {page_num+1}")
            
            # Third pass: Apply highlights based on positions found
            for field, position in field_positions.items():
                page_num = position['page']
                pos = position['pos']
                field_type = position['type']
                section = position['data']
                page = doc[page_num]
                
                # Get data from section
                issue = section.get('issue', '')
                regulation = section.get('regulation', '')
                recommendation = section.get('recommendation', '')
                
                # Try to find exact location with search
                field_instances = page.search_for(field)
                
                if field_instances:
                    for inst in field_instances:
                        if field_type == 'missing_field':
                            # Create a red rectangle around the field
                            field_rect = fitz.Rect(inst.x0, inst.y0, inst.x1 + 200, inst.y1 + 3)
                            rect_annot = page.add_rect_annot(field_rect)
                            rect_annot.set_colors(stroke=(1, 0, 0))  # Red stroke
                            rect_annot.set_border(width=1.5)  # Thicker border
                            rect_annot.update()
                            
                            # Also put a red highlight on the field name itself
                            highlight = page.add_highlight_annot(inst)
                            highlight.set_colors({"stroke": (1, 0, 0)})  # Red stroke
                            highlight.set_opacity(0.3)
                            highlight.update()
                            
                            # Add a text annotation beside it
                            point = fitz.Point(inst.x1 + 10, inst.y0)
                            note = page.add_text_annot(point, f"MISSING: Required information")
                            note.set_colors(stroke=(1, 0, 0))  # Red for text annotations
                            note.set_info({
                                "title": "Missing Information", 
                                "content": f"Issue: {issue}\nViolates: {regulation}\nRecommendation: {recommendation}"
                            })
                            note.update()
                            highlight_count += 3
                            
                        elif field_type == 'non_compliant_text':
                            # Add a highlight to the non-compliant text
                            highlight = page.add_highlight_annot(inst)
                            highlight.set_colors({"stroke": (1, 0, 0)})  # Red stroke
                            highlight.set_opacity(0.4)
                            
                            # Add annotation info
                            info = {
                                "title": "Non-Compliant Text",
                                "content": f"Issue: {issue}\nViolates: {regulation}\nRecommendation: {recommendation}"
                            }
                            highlight.set_info(info)
                            highlight.update()
                            highlight_count += 1
                else:
                    # If we couldn't find the exact position, try a more general approach
                    point = fitz.Point(100, pos + 30)
                    if field_type == 'missing_field':
                        msg = f"MISSING: {field}"
                    else:
                        msg = f"NON-COMPLIANT: {field[:30]}..."
                        
                    note = page.add_text_annot(point, msg)
                    note.set_colors(stroke=(1, 0, 0))
                    note.set_info({
                        "title": "Non-Compliant Section", 
                        "content": f"Issue: {issue}\nViolates: {regulation}\nRecommendation: {recommendation}"
                    })
                    note.update()
                    highlight_count += 1
            
            # Save the modified PDF - handle permission and encryption issues
            try:
                # First try with incremental update (faster, but may fail with encrypted PDFs)
                try:
                    doc.save(highlighted_path, incremental=True)
                    doc.close()
                    print(f"Created highlighted PDF at {highlighted_path} with {highlight_count} annotations")
                    return highlighted_path
                except Exception as e:
                    if "encryption" in str(e).lower():
                        # If encryption error, try saving to a new file without incremental update
                        print(f"PDF appears to be encrypted. Trying alternate save method...")
                        temp_path = highlighted_path + ".tmp"
                        doc.save(temp_path, incremental=False)
                        doc.close()
                        
                        # Replace the original file with the temporary file
                        if os.path.exists(highlighted_path):
                            try:
                                os.remove(highlighted_path)
                            except PermissionError:
                                print(f"Cannot remove existing file due to permissions. Using alternative filename.")
                                highlighted_path = highlighted_path.replace(".pdf", "_highlighted_new.pdf")
                        
                        os.rename(temp_path, highlighted_path)
                        print(f"Created highlighted PDF at {highlighted_path} with {highlight_count} annotations")
                        return highlighted_path
                    else:
                        # For other errors, re-raise
                        raise e
                        
            except Exception as save_error:
                print(f"Error saving PDF: {save_error}")
                try:
                    doc.close()
                except:
                    pass
                
                # If all save methods fail, create a separate summary document
                print("Creating a summary document instead")
                summary_path = highlighted_path.replace(".pdf", "_summary.pdf")
                
                try:
                    # Create new document with summary
                    doc = fitz.open()
                    new_page = doc.new_page()
                    new_page.insert_text((50, 50), "NON-COMPLIANT SECTIONS", fontsize=14, color=(1,0,0))
                    new_page.insert_text((50, 80), "The following sections in this document are non-compliant:", fontsize=12)
                    
                    y_pos = 110
                    for i, section in enumerate(non_compliant_sections, 1):
                        section_header = section.get('section_header', 'Unknown Section')
                        text = section.get('text', 'N/A')
                        regulation = section.get('regulation', 'N/A')
                        issue = section.get('issue', 'N/A')
                        recommendation = section.get('recommendation', 'N/A')
                        
                        new_page.insert_text((50, y_pos), f"{i}. Section: {section_header}", fontsize=11, color=(1,0,0))
                        y_pos += 20
                        new_page.insert_text((70, y_pos), f"Issue: {issue}", fontsize=10)
                        y_pos += 20
                        new_page.insert_text((70, y_pos), f"Problematic text/missing info: {text}", fontsize=10)
                        y_pos += 20
                        new_page.insert_text((70, y_pos), f"Violates: {regulation}", fontsize=10)
                        y_pos += 20
                        new_page.insert_text((70, y_pos), f"Recommendation: {recommendation}", fontsize=10)
                        y_pos += 30
                    
                    doc.save(summary_path)
                    doc.close()
                    print(f"Created summary document at {summary_path}")
                    return summary_path
                except Exception as e:
                    print(f"Error creating summary document: {e}")
                    return None
                
        except Exception as e:
            print(f"Error highlighting PDF: {str(e)}")
            import traceback
            traceback.print_exc()
            return None

    def _highlight_docx(self, original_path, highlighted_path, non_compliant_sections):
        """Highlight non-compliant sections in a DOCX document"""
        try:
            # Open the DOCX
            doc = docx.Document(original_path)
            
            # Find empty fields and non-compliant text
            empty_fields = {}
            for section in non_compliant_sections:
                issue = section.get('issue', '').lower()
                text = section.get('text', '').lower()
                section_header = section.get('section_header', '')
                
                # Check if this is a missing field issue
                if ('missing' in issue or 'empty' in issue or 'required' in issue or 
                    'missing' in text or 'empty' in text or 'not provided' in text):
                    empty_fields[section_header] = section
            
            # For each paragraph, check if it contains any non-compliant text or empty fields
            for paragraph in doc.paragraphs:
                original_text = paragraph.text
                
                # Highlight non-compliant text
                modified = False
                for section in non_compliant_sections:
                    non_compliant_text = section.get('text', '')
                    section_header = section.get('section_header', '')
                    
                    # Skip empty fields as they will be handled separately
                    if section_header in empty_fields:
                        continue
                    
                    # Only highlight actual text, not "missing" descriptions
                    if non_compliant_text and 'missing' not in non_compliant_text.lower() and non_compliant_text in original_text:
                        # Clear the paragraph and rebuild it with highlighted text
                        paragraph.clear()
                        
                        # Split the paragraph around the non-compliant text
                        parts = original_text.split(non_compliant_text)
                        
                        # Add back with highlighting
                        for i, part in enumerate(parts):
                            # Add normal text
                            if part:
                                run = paragraph.add_run(part)
                            
                            # Add highlighted text (except after the last part)
                            if i < len(parts) - 1:
                                run = paragraph.add_run(non_compliant_text)
                                run.font.color.rgb = RGBColor(255, 0, 0)  # Red color
                        
                        modified = True
                        break  # Break to avoid processing the same paragraph multiple times
                
                # If paragraph wasn't modified for non-compliant text, check for empty fields
                if not modified:
                    for field_name, section in empty_fields.items():
                        if field_name and field_name in original_text:
                            # Check if the field appears to be empty (just the label)
                            if len(original_text.strip()) <= len(field_name) + 2:  # Field name + possible colon and space
                                paragraph.clear()
                                
                                # Add the field name in red to indicate it's empty
                                run = paragraph.add_run(original_text)
                                run.font.color.rgb = RGBColor(255, 0, 0)  # Red color
                                run.bold = True
                                
                                # Add a note about the missing field
                                run = paragraph.add_run(" [MISSING REQUIRED INFORMATION]")
                                run.font.color.rgb = RGBColor(255, 0, 0)  # Red color
                                run.bold = True
                                break
            
            # Save the highlighted document
            doc.save(highlighted_path)
            print(f"Created highlighted DOCX at {highlighted_path}")
        except Exception as e:
            print(f"Error highlighting DOCX: {str(e)}")
            import traceback
            traceback.print_exc()

    def _highlight_txt(self, original_path, highlighted_path, non_compliant_sections):
        """Create an HTML version of the text file with highlighted sections"""
        try:
            # Read the original text file
            with open(original_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Find empty fields and non-compliant text
            empty_fields = {}
            for section in non_compliant_sections:
                issue = section.get('issue', '').lower()
                text = section.get('text', '').lower()
                section_header = section.get('section_header', '')
                
                # Check if this is a missing field issue
                if ('missing' in issue or 'empty' in issue or 'required' in issue or 
                    'missing' in text or 'empty' in text or 'not provided' in text):
                    empty_fields[section_header] = section
            
            # Create HTML with highlighting
            html_content = content
            
            # Escape HTML characters
            html_content = html.escape(html_content)
            
            # Replace each non-compliant section with a highlighted version
            for section in non_compliant_sections:
                text_to_highlight = html.escape(section.get('text', ''))
                section_header = section.get('section_header', '')
                
                # Skip empty fields as they will be handled separately
                if section_header in empty_fields:
                    continue
                
                # Only highlight actual text, not "missing" descriptions
                if text_to_highlight and 'missing' not in text_to_highlight.lower() and text_to_highlight in html_content:
                    highlighted_text = f'<span style="background-color: #ffcccc; color: red;">{text_to_highlight}</span>'
                    html_content = html_content.replace(text_to_highlight, highlighted_text)
            
            # Highlight empty fields
            for field_name, section in empty_fields.items():
                if field_name and field_name in html_content:
                    # Find the line containing the field
                    lines = html_content.split('\n')
                    for i, line in enumerate(lines):
                        if field_name in line:
                            # Check if the field appears to be empty
                            if len(line.strip()) <= len(field_name) + 2:  # Field name + possible colon and space
                                # Highlight the entire line
                                lines[i] = f'<span style="color: red; font-weight: bold;">{line} [MISSING REQUIRED INFORMATION]</span>'
                    
                    # Rejoin the lines
                    html_content = '\n'.join(lines)
            
            # Wrap in basic HTML structure
            html_content = f"""<!DOCTYPE html>
<html>
<head>
    <title>Highlighted Non-Compliant Text</title>
    <style>
        body {{ font-family: Arial, sans-serif; line-height: 1.5; }}
        .document {{ padding: 20px; }}
        h1 {{ color: #333; }}
        .footer {{ margin-top: 30px; font-size: 0.8em; color: #666; }}
    </style>
</head>
<body>
    <div class="document">
        <h1>Document with Highlighted Non-Compliant Text</h1>
        <pre>{html_content}</pre>
    </div>
    <div class="footer">
        <p>This document highlights sections that may not comply with relevant regulations.</p>
        <p>Red text indicates missing required information or non-compliant content.</p>
    </div>
</body>
</html>"""
            
            # Save as HTML
            html_path = highlighted_path.replace('.txt', '.html')
            with open(html_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            # Also save a copy with highlighting markers in txt format
            with open(highlighted_path, 'w', encoding='utf-8') as f:
                marked_content = content
                
                # Mark non-compliant text
                for section in non_compliant_sections:
                    text = section.get('text', '')
                    section_header = section.get('section_header', '')
                    
                    # Skip empty fields as they will be handled separately
                    if section_header in empty_fields:
                        continue
                    
                    # Only mark actual text, not "missing" descriptions
                    if text and 'missing' not in text.lower() and text in marked_content:
                        marked_content = marked_content.replace(text, f">>>RED_START>>>{text}<<<RED_END<<<")
                
                # Mark empty fields
                for field_name in empty_fields:
                    if field_name and field_name in marked_content:
                        # Find the line containing the field
                        lines = marked_content.split('\n')
                        for i, line in enumerate(lines):
                            if field_name in line:
                                # Check if the field appears to be empty
                                if len(line.strip()) <= len(field_name) + 2:  # Field name + possible colon and space
                                    # Mark the entire line
                                    lines[i] = f">>>RED_START>>>{line} [MISSING REQUIRED INFORMATION]<<<RED_END<<<"
                        
                        # Rejoin the lines
                        marked_content = '\n'.join(lines)
                
                f.write(marked_content)
                
            print(f"Created highlighted HTML at {html_path}")
            print(f"Created text version with markers at {highlighted_path}")
        except Exception as e:
            print(f"Error creating highlighted text file: {str(e)}")
            import traceback
            traceback.print_exc()

def kickoff():
    print(" starting the crew ...")
    ComplianceFlow().kickoff()
    print("\n Compliance Check Complete ...")

def plot():
    flow = ComplianceFlow()
    flow.plot()

if __name__ == "__main__":
    kickoff()